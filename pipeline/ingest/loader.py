from typing import Optional
import base64
import io
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from fastapi import HTTPException


@dataclass
class DocumentContent:
    raw_text: str
    pages: int
    format: str
    filename: str
    images: dict = field(default_factory=dict)


def _table_to_html(table) -> str:
    rows = table.extract()
    if not rows:
        return ""
    # Skip single-column tables — likely false positives
    if len(rows[0]) < 2:
        return ""

    def clean(cell) -> str:
        return (cell or "").strip()

    lines = ["<table>", "<thead><tr>"]
    for cell in rows[0]:
        lines.append(f"<th>{clean(cell)}</th>")
    lines.append("</tr></thead>")
    if len(rows) > 1:
        lines.append("<tbody>")
        for row in rows[1:]:
            lines.append("<tr>")
            for cell in row:
                lines.append(f"<td>{clean(cell)}</td>")
            lines.append("</tr>")
        lines.append("</tbody>")
    lines.append("</table>")
    return "".join(lines)


def load_pdf(data: bytes, filename: str, job_id: Optional[str] = None, uploads_dir: Optional[Path] = None) -> DocumentContent:
    doc = fitz.open(stream=data, filetype="pdf")
    pages_text = []
    for page in doc:
        # Find tables and record their bounding boxes
        tables = page.find_tables()
        table_entries = []  # list of (y0, html_string)
        table_rects = []    # fitz.Rect for each table to exclude overlapping text blocks
        for tbl in tables:
            html = _table_to_html(tbl)
            if html:
                bbox = tbl.bbox  # (x0, y0, x1, y1)
                table_entries.append((bbox[1], html))
                table_rects.append(fitz.Rect(bbox))

        # Extract text blocks, skipping any that overlap a table rect
        blocks = page.get_text("blocks")
        blocks.sort(key=lambda b: (round(b[1] / 20) * 20, b[0]))
        text_entries = []  # list of (y0, text)
        for b in blocks:
            if not b[4].strip() or b[6] != 0:
                continue
            block_rect = fitz.Rect(b[0], b[1], b[2], b[3])
            overlaps_table = any(block_rect.intersects(tr) for tr in table_rects)
            if not overlaps_table:
                text_entries.append((b[1], b[4].strip()))

        # Merge and sort by y-position
        all_entries = table_entries + text_entries
        all_entries.sort(key=lambda e: e[0])

        page_text = "\n\n".join(content for _, content in all_entries)
        pages_text.append(page_text)

    image_map: dict = {}
    if job_id and uploads_dir:
        img_dir = uploads_dir
        figure_counter = 1
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.width < 100 or pix.height < 100:
                        continue
                    if pix.n - pix.alpha > 3:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    img_filename = f"figure_{figure_counter}.png"
                    img_path = img_dir / img_filename
                    img_dir.mkdir(parents=True, exist_ok=True)
                    pix.save(str(img_path))
                    image_map[f"Figure {figure_counter}"] = {"url": f"/uploads/{job_id}/{img_filename}", "page": page_num}
                    figure_counter += 1
                except Exception:
                    continue

    doc.close()
    return DocumentContent(
        raw_text="\n".join(pages_text),
        pages=len(pages_text),
        format="pdf",
        filename=filename,
        images=image_map,
    )


def load_docx(data: bytes, filename: str) -> DocumentContent:
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return DocumentContent(
        raw_text="\n".join(parts),
        pages=len(doc.paragraphs),
        format="docx",
        filename=filename,
    )


def load_txt(data: bytes, filename: str) -> DocumentContent:
    text = data.decode("utf-8", errors="replace")
    lines = text.splitlines()
    return DocumentContent(
        raw_text=text,
        pages=len(lines),
        format="txt",
        filename=filename,
    )


def load_document(data: bytes, filename: str, job_id: Optional[str] = None, uploads_dir: Optional[Path] = None) -> DocumentContent:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return load_pdf(data, filename, job_id=job_id, uploads_dir=uploads_dir)
    elif lower.endswith(".docx"):
        return load_docx(data, filename)
    elif lower.endswith(".txt"):
        return load_txt(data, filename)
    else:
        raise HTTPException(status_code=415, detail=f"Unsupported file format: {filename}")
